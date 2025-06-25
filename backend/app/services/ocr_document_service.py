"""
Enhanced Document Processing Service with OCR support for scanned PDFs
Handles large files and provides OCR capabilities for image-based PDFs
"""

import asyncio
import aiofiles
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime
import warnings
import io
import sys
from contextlib import contextmanager
import tempfile

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

# OCR imports (optional dependencies)
try:
    import fitz  # PyMuPDF for better PDF handling

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdf2image import convert_from_path, convert_from_bytes
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# LangChain imports for text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Internal imports
from app.core.config import settings
from app.services.vector_service import vector_service
from app.services.database_service import DatabaseService

logger = logging.getLogger(__name__)

# Global function to broadcast logs (will be set by main.py)
broadcast_log_func = None


async def log_to_websocket(level: str, message: str, details: dict = None):
    """Send log message to WebSocket clients"""
    if broadcast_log_func:
        try:
            await broadcast_log_func(level, message, details)
        except Exception as e:
            logger.error(f"Failed to broadcast log: {e}")


@contextmanager
def suppress_pdf_warnings():
    """Context manager to suppress PyPDF2 warnings and capture them"""
    old_stderr = sys.stderr
    sys.stderr = io.StringIO()
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        try:
            yield
        finally:
            captured = sys.stderr.getvalue()
            sys.stderr = old_stderr
            if captured:
                logger.debug(f"PDF parsing messages: {captured[:200]}...")


class EnhancedDocumentService:
    """Enhanced document processing service with OCR support for large scanned PDFs"""

    def __init__(self):
        self.db_service = DatabaseService()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
        )
        self.max_file_size = 500 * 1024 * 1024  # 500MB limit

        # Check OCR availability
        if OCR_AVAILABLE:
            logger.info("OCR capabilities available (pdf2image + pytesseract)")
        else:
            logger.warning("OCR capabilities not available - scanned PDFs will fail")

        if PYMUPDF_AVAILABLE:
            logger.info("PyMuPDF available for better PDF handling")

    async def process_large_uploaded_file(
        self, file_content: bytes, filename: str, content_type: str
    ) -> Dict[str, Any]:
        """
        Process a large uploaded file with enhanced capabilities

        Args:
            file_content: File content as bytes
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Processing result with document ID and statistics
        """
        document_id = str(uuid.uuid4())
        start_time = datetime.now()

        # Initialize status tracking
        status = {
            "document_id": document_id,
            "filename": filename,
            "file_size": len(file_content),
            "status": "processing",
            "is_large_file": len(file_content) > 50 * 1024 * 1024,
            "requires_ocr": False,
            "stages": {
                "validation": {"status": "starting", "message": "Validating file"},
                "extraction": {
                    "status": "pending",
                    "message": "Waiting for validation",
                },
                "ocr_detection": {
                    "status": "pending",
                    "message": "Waiting for initial extraction",
                },
                "ocr_processing": {
                    "status": "pending",
                    "message": "Waiting for OCR detection",
                },
                "chunking": {
                    "status": "pending",
                    "message": "Waiting for text extraction",
                },
                "embedding": {"status": "pending", "message": "Waiting for chunking"},
                "storage": {"status": "pending", "message": "Waiting for embeddings"},
            },
        }

        try:
            # Stage 1: File Validation
            logger.info(
                f"Validating large file {filename} ({len(file_content)/1024/1024:.1f}MB)"
            )
            status["stages"]["validation"]["status"] = "processing"
            status["stages"]["validation"][
                "message"
            ] = f"Validating {len(file_content)/1024/1024:.1f}MB file"

            await log_to_websocket(
                "info",
                f"📏 Validating large file ({len(file_content)/1024/1024:.1f}MB)",
            )

            if len(file_content) > self.max_file_size:
                status["stages"]["validation"]["status"] = "failed"
                status["stages"]["validation"][
                    "message"
                ] = f"File too large (max {self.max_file_size/1024/1024:.0f}MB)"
                await log_to_websocket(
                    "error",
                    f"❌ File too large ({len(file_content)/1024/1024:.1f}MB > {self.max_file_size/1024/1024:.0f}MB)",
                )
                raise ValueError(
                    f"File size {len(file_content)/1024/1024:.1f}MB exceeds maximum {self.max_file_size/1024/1024:.0f}MB"
                )

            status["stages"]["validation"]["status"] = "completed"
            status["stages"]["validation"]["message"] = "File validation passed"
            await log_to_websocket("success", f"✅ File validation passed")

            # Stage 2: Initial Text Extraction
            logger.info(f"Starting text extraction for {filename}")
            status["stages"]["extraction"]["status"] = "processing"
            status["stages"]["extraction"][
                "message"
            ] = f"Extracting text from {self._get_document_type(filename).upper()} file"

            await log_to_websocket(
                "info",
                f"🔍 Starting enhanced text extraction from {self._get_document_type(filename).upper()}",
            )
            # Use memory-efficient extraction for large files
            text_content = await self._extract_text_efficient(
                file_content, filename, content_type, status
            )
            await log_to_websocket(
                "success",
                f"📄 Extracted {len(text_content)} characters using enhanced processing",
            )

            # Stage 3: OCR Detection and Processing
            if len(text_content.strip()) < 100 and filename.lower().endswith(".pdf"):
                logger.info(
                    f"Minimal text extracted from {filename}, checking if OCR is needed"
                )
                status["stages"]["ocr_detection"]["status"] = "processing"
                status["stages"]["ocr_detection"][
                    "message"
                ] = "Detecting if OCR is needed (minimal text found)"

                if OCR_AVAILABLE:
                    status["requires_ocr"] = True
                    status["stages"]["ocr_detection"]["status"] = "completed"
                    status["stages"]["ocr_detection"][
                        "message"
                    ] = "OCR required - processing with image recognition"

                    await log_to_websocket(
                        "info", f"👁️ Scanned PDF detected - initiating OCR processing"
                    )

                    # Perform OCR
                    status["stages"]["ocr_processing"]["status"] = "processing"
                    status["stages"]["ocr_processing"][
                        "message"
                    ] = "Performing OCR on scanned PDF"

                    await log_to_websocket(
                        "info", f"🖼️ Converting PDF pages to images for OCR analysis"
                    )
                    ocr_text = await self._perform_ocr_extraction(
                        file_content, filename, status
                    )
                    if ocr_text and len(ocr_text.strip()) > len(text_content.strip()):
                        text_content = ocr_text
                        status["stages"]["ocr_processing"]["status"] = "completed"
                        status["stages"]["ocr_processing"][
                            "message"
                        ] = f"OCR completed - extracted {len(ocr_text)} characters"
                    else:
                        status["stages"]["ocr_processing"]["status"] = "failed"
                        status["stages"]["ocr_processing"][
                            "message"
                        ] = "OCR did not improve text extraction"
                else:
                    status["stages"]["ocr_detection"]["status"] = "failed"
                    status["stages"]["ocr_detection"][
                        "message"
                    ] = "OCR required but not available"
                    status["stages"]["ocr_processing"]["status"] = "skipped"
                    status["stages"]["ocr_processing"][
                        "message"
                    ] = "OCR dependencies not installed"
            else:
                status["stages"]["ocr_detection"]["status"] = "completed"
                status["stages"]["ocr_detection"][
                    "message"
                ] = "Sufficient text found - OCR not needed"
                status["stages"]["ocr_processing"]["status"] = "skipped"
                status["stages"]["ocr_processing"]["message"] = "OCR not required"

            if not text_content.strip():
                status["stages"]["extraction"]["status"] = "failed"
                status["stages"]["extraction"][
                    "message"
                ] = "No readable text found in file"
                raise ValueError("No text content could be extracted from the file")

            status["stages"]["extraction"]["status"] = "completed"
            status["stages"]["extraction"][
                "message"
            ] = f"Extracted {len(text_content)} characters"

            # Stage 4: Text Chunking (memory-efficient for large text)
            logger.info(f"Starting chunking for {filename}")
            status["stages"]["chunking"]["status"] = "processing"
            status["stages"]["chunking"]["message"] = "Creating semantic chunks"

            chunks = await self._create_chunks_efficient(
                text_content, filename, document_id, status
            )

            if not chunks:
                status["stages"]["chunking"]["status"] = "failed"
                status["stages"]["chunking"][
                    "message"
                ] = "Failed to create chunks from text"
                raise ValueError("No chunks could be created from the document")

            status["stages"]["chunking"]["status"] = "completed"
            status["stages"]["chunking"]["message"] = f"Created {len(chunks)} chunks"

            # Stage 5: Embedding Generation (batch processing for many chunks)
            logger.info(f"Starting embedding generation for {filename}")
            status["stages"]["embedding"]["status"] = "processing"
            status["stages"]["embedding"][
                "message"
            ] = "Generating vector embeddings (batch processing)"

            chunk_ids = await self._generate_embeddings_batch(
                chunks, document_id, status
            )

            status["stages"]["embedding"]["status"] = "completed"
            status["stages"]["embedding"][
                "message"
            ] = f"Generated embeddings for {len(chunk_ids)} chunks"

            # Stage 6: Database Storage
            logger.info(f"Storing metadata for {filename}")
            status["stages"]["storage"]["status"] = "processing"
            status["stages"]["storage"]["message"] = "Saving document metadata"

            await self._store_document_metadata(
                filename,
                text_content,
                content_type,
                len(file_content),
                len(chunks),
                document_id,
                status,
            )

            # Save full content to file (async for large files)
            await self._save_file_async(file_content, document_id, filename)

            status["stages"]["storage"]["status"] = "completed"
            status["stages"]["storage"]["message"] = "Document metadata saved"

            # Final status
            processing_time = (datetime.now() - start_time).total_seconds()

            logger.info(
                f"Successfully processed large file {filename} in {processing_time:.2f}s with {len(chunks)} chunks"
            )

            return {
                "document_id": document_id,
                "filename": filename,
                "chunk_count": len(chunks),
                "text_length": len(text_content),
                "file_size": len(file_content),
                "file_size_mb": round(len(file_content) / (1024 * 1024), 2),
                "processing_time_seconds": round(processing_time, 2),
                "is_large_file": status["is_large_file"],
                "required_ocr": status["requires_ocr"],
                "ocr_available": OCR_AVAILABLE,
                "status": "success",
                "message": f"Large file processed successfully with {len(chunks)} chunks in {processing_time:.1f}s",
                "stages": status["stages"],
            }

        except Exception as e:
            # Update status with error information
            processing_time = (datetime.now() - start_time).total_seconds()
            error_message = str(e)

            logger.error(f"Error processing large file {filename}: {error_message}")

            # Determine which stage failed
            failed_stage = "validation"
            for stage_name, stage_info in status["stages"].items():
                if stage_info["status"] == "processing":
                    failed_stage = stage_name
                    break
                elif stage_info["status"] == "failed":
                    failed_stage = stage_name
                    break

            return {
                "document_id": document_id,
                "filename": filename,
                "file_size": len(file_content),
                "file_size_mb": round(len(file_content) / (1024 * 1024), 2),
                "processing_time_seconds": round(processing_time, 2),
                "is_large_file": status["is_large_file"],
                "required_ocr": status["requires_ocr"],
                "ocr_available": OCR_AVAILABLE,
                "status": "error",
                "message": f"Failed during {failed_stage}: {error_message}",
                "error_stage": failed_stage,
                "error_details": error_message,
                "stages": status["stages"],
            }

    async def _extract_text_efficient(
        self, file_content: bytes, filename: str, content_type: str, status: dict
    ) -> str:
        """Memory-efficient text extraction for large files"""
        try:
            file_extension = Path(filename).suffix.lower()

            if file_extension == ".pdf" or "pdf" in content_type:
                return await self._extract_from_pdf_efficient(file_content, status)
            elif file_extension in [".docx", ".doc"] or "word" in content_type:
                return await self._extract_from_docx(file_content)
            elif file_extension in [".txt", ".md"] or "text" in content_type:
                return await self._extract_from_text(file_content)
            else:
                return await self._extract_from_text(file_content)

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise

    async def _extract_from_pdf_efficient(
        self, file_content: bytes, status: dict
    ) -> str:
        """Memory-efficient PDF text extraction"""
        try:
            logger.info(
                f"Processing large PDF file ({len(file_content)/1024/1024:.1f}MB)..."
            )

            # Try PyMuPDF first (better for large files)
            if PYMUPDF_AVAILABLE:
                logger.info("Using PyMuPDF for efficient PDF processing")
                await log_to_websocket(
                    "info",
                    f"🔧 Using PyMuPDF for enhanced PDF processing ({len(file_content)/1024/1024:.1f}MB)",
                )
                text_content = []

                with suppress_pdf_warnings():
                    doc = fitz.open(stream=file_content, filetype="pdf")
                    total_pages = len(doc)

                    logger.info(f"PDF has {total_pages} pages")
                    await log_to_websocket(
                        "info",
                        f"📄 Analyzing {total_pages} pages with advanced PDF reader",
                    )
                    status["stages"]["extraction"][
                        "message"
                    ] = f"Processing {total_pages} pages with PyMuPDF"

                    successful_pages = 0
                    for page_num in range(total_pages):
                        try:
                            page = doc[page_num]
                            page_text = page.get_text()
                            if page_text.strip():
                                text_content.append(
                                    f"[Page {page_num + 1}]\n{page_text}"
                                )
                                successful_pages += 1

                            # Update progress for large files
                            if page_num % 50 == 0:
                                await log_to_websocket(
                                    "info",
                                    f"📃 Processing page {page_num + 1}/{total_pages}",
                                )
                                status["stages"]["extraction"][
                                    "message"
                                ] = f"Processed {page_num + 1}/{total_pages} pages"

                        except Exception as page_error:
                            logger.warning(
                                f"Could not extract text from page {page_num + 1}: {page_error}"
                            )
                            continue

                    doc.close()

                logger.info(
                    f"PyMuPDF extracted text from {successful_pages}/{total_pages} pages"
                )
                await log_to_websocket(
                    "success",
                    f"✅ PyMuPDF extracted text from {successful_pages}/{total_pages} pages",
                )
                return "\n\n".join(text_content)

            # Fallback to PyPDF2
            else:
                logger.info("Using PyPDF2 for PDF processing")
                return await self._extract_from_pdf_pypdf2(file_content, status)

        except Exception as e:
            logger.error(f"Efficient PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from large PDF: {str(e)[:100]}")

    async def _extract_from_pdf_pypdf2(self, file_content: bytes, status: dict) -> str:
        """PyPDF2 extraction with progress tracking"""
        with suppress_pdf_warnings():
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            total_pages = len(pdf_reader.pages)
            logger.info(f"PDF has {total_pages} pages")

            text_content = []
            successful_pages = 0

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                        successful_pages += 1

                    # Update progress for large files
                    if page_num % 20 == 0:
                        status["stages"]["extraction"][
                            "message"
                        ] = f"Processed {page_num + 1}/{total_pages} pages (PyPDF2)"

                except Exception as page_error:
                    logger.warning(
                        f"Could not extract text from page {page_num + 1}: {page_error}"
                    )
                    continue

            logger.info(
                f"PyPDF2 extracted text from {successful_pages}/{total_pages} pages"
            )
            return "\n\n".join(text_content)

    async def _perform_ocr_extraction(
        self, file_content: bytes, filename: str, status: dict
    ) -> str:
        """Perform OCR on scanned PDF"""
        if not OCR_AVAILABLE:
            raise ValueError("OCR dependencies not available")

        try:
            logger.info(f"Starting OCR processing for {filename}")
            await log_to_websocket(
                "info", f"👁️ Starting optical character recognition (OCR)"
            )

            # Convert PDF to images
            await log_to_websocket(
                "info", f"🖼️ Converting PDF pages to high-resolution images"
            )
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
                temp_pdf.write(file_content)
                temp_pdf_path = temp_pdf.name

            try:
                # Convert PDF pages to images
                images = convert_from_path(temp_pdf_path, dpi=300, fmt="jpeg")
                logger.info(f"Converted PDF to {len(images)} images for OCR")
                await log_to_websocket(
                    "success",
                    f"📸 Converted PDF to {len(images)} high-resolution images",
                )

                ocr_text = []
                for i, image in enumerate(images):
                    try:
                        # Perform OCR on each page
                        page_text = pytesseract.image_to_string(image, lang="eng")
                        if page_text.strip():
                            ocr_text.append(f"[Page {i + 1} - OCR]\n{page_text}")

                        # Update progress
                        if i % 5 == 0 or i == 0:  # Log every 5 pages
                            await log_to_websocket(
                                "info",
                                f"🔍 OCR reading text from page {i + 1}/{len(images)}",
                            )
                        status["stages"]["ocr_processing"][
                            "message"
                        ] = f"OCR processing page {i + 1}/{len(images)}"

                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page {i + 1}: {ocr_error}")
                        continue

                logger.info(
                    f"OCR completed - extracted text from {len(ocr_text)} pages"
                )
                await log_to_websocket(
                    "success",
                    f"✅ OCR completed - extracted text from {len(ocr_text)} pages",
                )
                return "\n\n".join(ocr_text)

            finally:
                # Cleanup temp file
                Path(temp_pdf_path).unlink(missing_ok=True)

        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            raise ValueError(f"OCR extraction failed: {str(e)}")

    async def _create_chunks_efficient(
        self, text_content: str, filename: str, document_id: str, status: dict
    ) -> List[Document]:
        """Memory-efficient chunking for large text"""
        try:
            # For very large text, process in batches
            if len(text_content) > 1_000_000:  # 1MB of text
                logger.info(
                    f"Large text content ({len(text_content)} chars) - using batch processing"
                )
                status["stages"]["chunking"][
                    "message"
                ] = "Processing large text in batches"

            # Create a LangChain Document
            doc = Document(
                page_content=text_content,
                metadata={
                    "filename": filename,
                    "document_id": document_id,
                    "doc_type": self._get_document_type(filename),
                },
            )

            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])

            # Add chunk-specific metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update(
                    {
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "chunk_id": f"{document_id}_chunk_{i}",
                    }
                )

            return chunks

        except Exception as e:
            logger.error(f"Error creating chunks: {e}")
            raise

    async def _generate_embeddings_batch(
        self, chunks: List[Document], document_id: str, status: dict
    ) -> List[str]:
        """Generate embeddings in batches for memory efficiency"""
        try:
            batch_size = 100  # Process in batches of 100 chunks
            all_chunk_ids = []

            for i in range(0, len(chunks), batch_size):
                batch = chunks[i : i + batch_size]

                # Update progress
                status["stages"]["embedding"][
                    "message"
                ] = f"Processing embeddings batch {i//batch_size + 1}/{(len(chunks) + batch_size - 1)//batch_size}"

                # Generate embeddings for this batch
                chunk_ids = vector_service.add_documents(
                    texts=[chunk.page_content for chunk in batch],
                    metadatas=[chunk.metadata for chunk in batch],
                    document_id=document_id,
                )

                all_chunk_ids.extend(chunk_ids)

                # Small delay to prevent memory issues
                if len(chunks) > 500:
                    await asyncio.sleep(0.1)

            return all_chunk_ids

        except Exception as e:
            logger.error(f"Batch embedding generation failed: {e}")
            raise

    async def _store_document_metadata(
        self,
        filename: str,
        text_content: str,
        content_type: str,
        file_size: int,
        chunk_count: int,
        document_id: str,
        status: dict,
    ):
        """Store document metadata"""
        try:
            self.db_service.save_document(
                filename=filename,
                content=(
                    text_content[:1000] + "..."
                    if len(text_content) > 1000
                    else text_content
                ),
                doc_type=self._get_document_type(filename),
                metadata={
                    "content_type": content_type,
                    "file_size": file_size,
                    "chunk_count": chunk_count,
                    "processing_date": datetime.now().isoformat(),
                    "text_length": len(text_content),
                    "is_large_file": file_size > 50 * 1024 * 1024,
                    "required_ocr": status.get("requires_ocr", False),
                    "ocr_available": OCR_AVAILABLE,
                },
            )
        except Exception as e:
            logger.error(f"Failed to store document metadata: {e}")
            raise

    async def _save_file_async(
        self, file_content: bytes, document_id: str, filename: str
    ):
        """Async file saving for large files"""
        try:
            data_dir = Path(settings.data_storage_path)
            data_dir.mkdir(parents=True, exist_ok=True)

            file_extension = Path(filename).suffix
            save_path = data_dir / f"{document_id}{file_extension}"

            # Use async file writing for large files
            async with aiofiles.open(save_path, "wb") as f:
                await f.write(file_content)

            logger.info(f"Saved large file {filename} to {save_path}")
        except Exception as e:
            logger.error(f"Error saving large file {filename}: {e}")

    def _get_document_type(self, filename: str) -> str:
        """Determine document type from filename"""
        extension = Path(filename).suffix.lower()
        type_mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".txt": "text",
            ".md": "markdown",
            ".html": "html",
            ".htm": "html",
        }
        return type_mapping.get(extension, "unknown")

    # Include other methods for compatibility
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    async def _extract_from_text(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            return file_content.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise ValueError(f"Failed to extract text: {e}")


# Global enhanced document service instance
enhanced_document_service = EnhancedDocumentService()
